"""
pdf_reader.py — Teklif Mukayese Yardımcısı
==========================================
Klasördeki tüm PDF, Excel, Word (.docx) ve Outlook (.msg) dosyalarını okur,
ham metni ekrana basar. Claude bu çıktıyı ana konuşmada analiz eder.

Kullanım:
    python pdf_reader.py "C:/Klasor/Yolu"

Gereksinimler (otomatik kontrol edilir, eksikse uyarı verir):
    pip install pdfplumber openpyxl python-docx extract-msg
"""

import sys
import os
import glob

# Windows terminal encoding sorunu: UTF-8 zorla
if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')


def main():
    if len(sys.argv) < 2:
        print("Kullanim: python pdf_reader.py \"KLASOR_YOLU\"")
        sys.exit(1)

    klasor = sys.argv[1]

    if not os.path.isdir(klasor):
        print(f"HATA: Klasor bulunamadi: {klasor}")
        sys.exit(1)

    # ── PDF'leri oku ────────────────────────────────────────────────────────
    pdf_dosyalari = sorted(glob.glob(os.path.join(klasor, "*.pdf")))
    if pdf_dosyalari:
        try:
            import pdfplumber
        except ImportError:
            print("HATA: pdfplumber yuklu degil. Yuklemek icin: pip install pdfplumber")
            sys.exit(1)

        for f in pdf_dosyalari:
            print(f"\n{'='*70}")
            print(f"PDF: {os.path.basename(f)}")
            print('='*70)
            try:
                with pdfplumber.open(f) as pdf:
                    metin_var = False
                    for i, page in enumerate(pdf.pages):
                        txt = page.extract_text()
                        if txt and txt.strip():
                            metin_var = True
                            print(f"\n--- Sayfa {i+1} ---")
                            print(txt)
                        tablolar = page.extract_tables()
                        if tablolar:
                            for j, tablo in enumerate(tablolar):
                                print(f"\n[TABLO {i+1}.{j+1}]")
                                for satir in tablo:
                                    print('\t'.join(str(h or '') for h in satir))
                    if not metin_var:
                        print("UYARI: Metin cikarilamadi — gorsel/taranmis PDF olabilir.")
                        print("       Read tool ile direkt acin (Claude multimodal gorebilir).")
            except Exception as e:
                print(f"HATA: {e}")
    else:
        print("(PDF dosyasi bulunamadi)")

    # ── Excel / XLS dosyalarını oku ─────────────────────────────────────────
    excel_dosyalari = sorted(
        glob.glob(os.path.join(klasor, "*.xlsx")) +
        glob.glob(os.path.join(klasor, "*.xls"))
    )
    if excel_dosyalari:
        try:
            import openpyxl
        except ImportError:
            print("HATA: openpyxl yuklu degil. Yuklemek icin: pip install openpyxl")
            sys.exit(1)

        for f in excel_dosyalari:
            print(f"\n{'='*70}")
            print(f"EXCEL: {os.path.basename(f)}")
            print('='*70)
            try:
                wb = openpyxl.load_workbook(f, data_only=True)
                for sn in wb.sheetnames:
                    ws = wb[sn]
                    print(f"\n-- Sheet: {sn} --")
                    for satir in ws.iter_rows(values_only=True):
                        if any(h is not None for h in satir):
                            print(satir)
            except Exception as e:
                print(f"HATA: {e}")
    else:
        print("(Excel dosyasi bulunamadi)")

    # ── Word (.docx) dosyalarını oku ────────────────────────────────────────
    docx_dosyalari = sorted(glob.glob(os.path.join(klasor, "*.docx")))
    if docx_dosyalari:
        docx_lib = None
        try:
            import docx as docx_lib
        except ImportError:
            print("UYARI: python-docx yuklu degil -> .docx dosyalari atlandi.")
            print("       Yuklemek icin: pip install python-docx")

        if docx_lib:
            for f in docx_dosyalari:
                print(f"\n{'='*70}")
                print(f"WORD: {os.path.basename(f)}")
                print('='*70)
                try:
                    doc = docx_lib.Document(f)
                    for para in doc.paragraphs:
                        if para.text.strip():
                            print(para.text)
                    for i, tablo in enumerate(doc.tables):
                        print(f"\n[TABLO {i+1}]")
                        for satir in tablo.rows:
                            hucreler = [c.text.strip() for c in satir.cells]
                            if any(hucreler):
                                print('\t'.join(hucreler))
                except Exception as e:
                    print(f"HATA: {e}")
    else:
        print("(Word dosyasi bulunamadi)")

    # ── Outlook MSG (.msg) dosyalarını oku ──────────────────────────────────
    # Tedarikçiler tekliflerini mail gövdesine yazıp .msg olarak kaydedilmişse
    # bu bölüm devreye girer. extract-msg kütüphanesi gereklidir.
    msg_dosyalari = sorted(glob.glob(os.path.join(klasor, "*.msg")))
    if msg_dosyalari:
        emsg = None
        try:
            import extract_msg as emsg
        except ImportError:
            print("UYARI: extract-msg yuklu degil -> .msg dosyalari atlandi.")
            print("       Yuklemek icin: python -m pip install extract-msg")

        if emsg:
            for f in msg_dosyalari:
                print(f"\n{'='*70}")
                print(f"OUTLOOK MSG: {os.path.basename(f)}")
                print('='*70)
                try:
                    msg = emsg.openMsg(f)
                    print(f"Kimden : {msg.sender}")
                    print(f"Konu   : {msg.subject}")
                    print(f"Tarih  : {msg.date}")
                    print("--- MAIL GOVDESI ---")
                    print(msg.body or "(govde bos)")
                except Exception as e:
                    print(f"HATA: {e}")
    else:
        print("(Outlook MSG dosyasi bulunamadi)")

    print(f"\n{'='*70}")
    print("OKUMA TAMAMLANDI — Yukardaki ciktiyi analiz edin.")
    print('='*70)


if __name__ == "__main__":
    main()
